"""build_docx.py — mythos-codex 블로그 docx 빌더"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).parent
FINAL_MD = BASE / "4.final.md"
IMAGES_DIR = BASE / "images"
OUTPUT = BASE / "5.Claude Mythos.docx"


def set_korean_font(run, kor="맑은 고딕", eng="Calibri", size_pt=11, bold=False, color_hex=None):
    run.font.name = eng
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), kor)
    r_fonts.set(qn("w:ascii"), eng)
    r_fonts.set(qn("w:hAnsi"), eng)


def set_default_korean(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "맑은 고딕")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def add_heading_styled(doc, text, level):
    heading = doc.add_heading(text, level=level)
    colors = {1: "1F4E79", 2: "2E75B6", 3: "5B9BD5"}
    sizes = {1: 22, 2: 16, 3: 13}
    for run in heading.runs:
        set_korean_font(
            run,
            size_pt=sizes.get(level, 12),
            bold=True,
            color_hex=colors.get(level, "1F4E79"),
        )


def add_paragraph_with_inline(doc, text):
    paragraph = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_korean_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_korean_font(run)
    return paragraph


def add_image_safe(doc, image_path, alt_text=""):
    path = Path(image_path)
    if not path.exists() or path.stat().st_size == 0:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[이미지 누락: {path.name}]")
        set_korean_font(run, size_pt=10, color_hex="999999")
        return
    doc.add_picture(str(path), width=Inches(5.9))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(alt_text)
        set_korean_font(run, size_pt=9, color_hex="595959")


def parse_and_build(doc, content):
    lines = content.splitlines()
    index = 0

    if lines and lines[0].strip() == "---":
        index = 1
        while index < len(lines) and lines[index].strip() != "---":
            index += 1
        index += 1

    while index < len(lines):
        line = lines[index]

        match = re.match(r"^(#{1,3})\s+(.*)", line)
        if match:
            add_heading_styled(doc, match.group(2).strip(), len(match.group(1)))
            index += 1
            continue

        match = re.match(r"^!\[([^\]]*)\]\(images/([^)]+)\)", line.strip())
        if match:
            add_image_safe(doc, IMAGES_DIR / match.group(2), match.group(1))
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        add_paragraph_with_inline(doc, line.rstrip())
        index += 1


def build():
    if not FINAL_MD.exists():
        raise FileNotFoundError(f"입력 파일 없음: {FINAL_MD}")

    doc = Document()
    set_default_korean(doc)
    content = FINAL_MD.read_text(encoding="utf-8")
    parse_and_build(doc, content)
    doc.save(str(OUTPUT))

    if not OUTPUT.exists() or OUTPUT.stat().st_size == 0:
        raise RuntimeError(f"검증 실패: {OUTPUT}")

    print(f"[OK] saved: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        sys.exit(1)
