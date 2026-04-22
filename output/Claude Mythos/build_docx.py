"""build_docx.py — Claude Mythos 블로그 docx 빌더"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
FINAL_MD = BASE / "4.final.md"
IMAGES_DIR = BASE / "images"
OUTPUT = BASE / "5.Claude Mythos.docx"


def set_korean_font(run, kor="맑은 고딕", eng="Calibri", size_pt=11, bold=False, color_hex=None):
    run.font.name = eng
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), kor)
    rFonts.set(qn("w:ascii"), eng)
    rFonts.set(qn("w:hAnsi"), eng)


def set_default_korean(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")


def add_image_safe(doc, img_path, alt_text=""):
    p = Path(img_path)
    if not p.exists() or p.stat().st_size == 0:
        para = doc.add_paragraph()
        run = para.add_run(f"[이미지: {p.name}]")
        set_korean_font(run, size_pt=10, color_hex="999999")
        return
    doc.add_picture(str(p), width=Inches(5.9))
    if alt_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(alt_text)
        set_korean_font(run, size_pt=9, color_hex="595959")


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    colors = {1: "1F4E79", 2: "2E75B6", 3: "5B9BD5"}
    sizes = {1: 22, 2: 16, 3: 13}
    color = colors.get(level, "1F4E79")
    size = sizes.get(level, 14)
    for run in h.runs:
        set_korean_font(run, size_pt=size, bold=True, color_hex=color)


def add_paragraph_with_inline(doc, text):
    """Handle **bold** inline markdown."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_korean_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_korean_font(run)
    return p


def parse_and_build(doc, content):
    lines = content.splitlines()
    i = 0
    # Skip YAML frontmatter
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1  # skip closing ---

    while i < len(lines):
        line = lines[i]

        # Skip HTML comments
        if line.strip().startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            add_heading_styled(doc, m.group(2).strip(), level)
            i += 1
            continue

        # Images: ![alt](images/filename.png)
        m = re.match(r'^!\[([^\]]*)\]\(images/([^)]+)\)', line.strip())
        if m:
            alt = m.group(1)
            fname = m.group(2)
            add_image_safe(doc, IMAGES_DIR / fname, alt)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph (strip trailing 2-space line breaks)
        text = line.rstrip()
        if text:
            add_paragraph_with_inline(doc, text)
        i += 1


def build():
    doc = Document()
    set_default_korean(doc)
    content = FINAL_MD.read_text(encoding="utf-8")
    parse_and_build(doc, content)
    doc.save(str(OUTPUT))
    if not OUTPUT.exists() or OUTPUT.stat().st_size == 0:
        raise RuntimeError(f"검증 실패: {OUTPUT}")
    print(f"[OK] saved: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(build())
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
